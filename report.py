import io
from datetime import datetime
from typing import List, Dict, Optional

from analysis import get_sieve_label, generate_export_data, generate_summary_export


def generate_word_report(samples_data: List[Dict], title: str = "火山灰筛分分析报告") -> bytes:
    """
    生成 Word 格式分析报告
    依赖 python-docx
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")
    
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    doc.add_heading("一、样本汇总", level=1)
    
    summary_df = generate_summary_export(samples_data)
    if not summary_df.empty:
        table = doc.add_table(rows=1, cols=len(summary_df.columns))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(summary_df.columns):
            hdr_cells[i].text = str(col)
            for para in hdr_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        
        for _, row in summary_df.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(summary_df.columns):
                val = row[col]
                if isinstance(val, float):
                    row_cells[i].text = f"{val:.3f}"
                else:
                    row_cells[i].text = str(val)
                for para in row_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    for idx, sample in enumerate(samples_data, 1):
        analysis = sample.get('analysis', {})
        
        doc.add_heading(f"二、样本 {sample['sample_no']} 详细分析", level=1)
        
        doc.add_heading("2.1 基本信息", level=2)
        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Table Grid'
        
        info_items = [
            ("样本编号", sample.get('sample_no', '')),
            ("采样点", sample.get('sampling_site', '')),
            ("喷发层位", sample.get('eruption_layer', '未记录')),
            ("分组", sample.get('group_name', '未分组')),
            ("采样时间", sample.get('sampling_time', '未记录')),
            ("深度", f"{sample.get('depth', '未记录')} m" if sample.get('depth') else '未记录'),
            ("剖面位置", sample.get('profile_position', '未记录')),
            ("样本总重量", f"{sample.get('total_weight', 0):.2f} g"),
            ("筛分总重量", f"{analysis.get('total_retained', 0):.2f} g"),
            ("底盘重量", f"{analysis.get('pan_weight', 0):.4f} g"),
            ("创建时间", sample.get('created_at', '')),
            ("备注", sample.get('description', '无')),
        ]
        
        for key, value in info_items:
            row_cells = info_table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
            row_cells[0].width = Cm(3)
            for para in row_cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True
        
        doc.add_paragraph()
        
        doc.add_heading("2.2 粒径参数", level=2)
        params_table = doc.add_table(rows=1, cols=2)
        params_table.style = 'Light Grid Accent 1'
        
        hdr = params_table.rows[0].cells
        hdr[0].text = '参数'
        hdr[1].text = '数值'
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        
        param_items = [
            ("D10", analysis.get('d10')),
            ("D25", analysis.get('d25')),
            ("中值粒径 D50", analysis.get('median_diameter')),
            ("D75", analysis.get('d75')),
            ("D90", analysis.get('d90')),
            ("分选系数", analysis.get('sorting_coefficient')),
            ("沉积分选评价", analysis.get('sorting_grade', '无法判断')),
            ("可信度评分", f"{analysis.get('reliability_score', 0)} / 100"),
            ("可信度等级", analysis.get('reliability_level', '未知')),
        ]
        
        for name, val in param_items:
            row = params_table.add_row().cells
            row[0].text = name
            if isinstance(val, float):
                row[1].text = f"{val:.3f} mm" if 'D' in name or '粒径' in name or '系数' in name else str(val)
            else:
                row[1].text = str(val) if val else '—'
        
        doc.add_paragraph()
        
        doc.add_heading("2.3 筛分数据明细", level=2)
        df = analysis.get('df')
        if df is not None and not df.empty:
            detail_table = doc.add_table(rows=1, cols=4)
            detail_table.style = 'Light Grid Accent 1'
            
            hdr = detail_table.rows[0].cells
            headers = ['粒级', '筛孔 (mm)', '留存重量 (g)', '占比 (%)']
            for i, h in enumerate(headers):
                hdr[i].text = h
                for para in hdr[i].paragraphs:
                    for run in para.runs:
                        run.bold = True
            
            for _, row_data in df.iterrows():
                row = detail_table.add_row().cells
                row[0].text = get_sieve_label(row_data['sieve_size'])
                row[1].text = f"{row_data['sieve_size']}"
                row[2].text = f"{row_data['retained_weight']:.4f}"
                row[3].text = f"{row_data['weight_percent']:.2f}"
            
            pan_w = analysis.get('pan_weight', 0)
            pan_p = analysis.get('pan_percent', 0)
            if pan_w > 0:
                row = detail_table.add_row().cells
                row[0].text = '底盘（小于最小粒级）'
                row[1].text = '—'
                row[2].text = f"{pan_w:.4f}"
                row[3].text = f"{pan_p:.2f}"
        
        if idx < len(samples_data):
            doc.add_page_break()
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_pdf_report(samples_data: List[Dict], title: str = "火山灰筛分分析报告") -> bytes:
    """
    生成 PDF 格式分析报告
    依赖 reportlab
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm, mm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        raise ImportError("请安装 reportlab: pip install reportlab")
    
    buffer = io.BytesIO()
    
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=2*cm,
        bottomMargin=2*cm,
        leftMargin=2*cm,
        rightMargin=2*cm,
    )
    
    styles = getSampleStyleSheet()
    
    chinese_font_found = False
    try:
        font_paths = [
            '/System/Library/Fonts/PingFang.ttc',
            '/System/Library/Fonts/STHeiti Light.ttc',
            '/Library/Fonts/Arial Unicode.ttf',
            '/System/Library/Fonts/Hiragino Sans GB.ttc',
        ]
        for font_path in font_paths:
            import os
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    chinese_font_found = True
                    break
                except Exception:
                    continue
    except Exception:
        pass
    
    font_name = 'ChineseFont' if chinese_font_found else 'Helvetica'
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontName=font_name,
        fontSize=18,
        spaceAfter=20,
        alignment=1,
    )
    h1_style = ParagraphStyle(
        'CustomH1',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=14,
        spaceBefore=15,
        spaceAfter=10,
    )
    h2_style = ParagraphStyle(
        'CustomH2',
        parent=styles['Heading2'],
        fontName=font_name,
        fontSize=12,
        spaceBefore=10,
        spaceAfter=6,
    )
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        spaceAfter=6,
    )
    info_style = ParagraphStyle(
        'CustomInfo',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=9,
        textColor=colors.grey,
        alignment=1,
        spaceAfter=20,
    )
    
    story = []
    
    story.append(Paragraph(title, title_style))
    story.append(Paragraph(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", info_style))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("一、样本汇总", h1_style))
    
    summary_df = generate_summary_export(samples_data)
    if not summary_df.empty:
        table_data = []
        table_data.append(list(summary_df.columns))
        
        for _, row in summary_df.iterrows():
            row_data = []
            for col in summary_df.columns:
                val = row[col]
                if isinstance(val, float):
                    row_data.append(f"{val:.3f}")
                else:
                    row_data.append(str(val))
            table_data.append(row_data)
        
        table = Table(table_data, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ]))
        story.append(table)
    
    story.append(PageBreak())
    
    for idx, sample in enumerate(samples_data, 1):
        analysis = sample.get('analysis', {})
        
        story.append(Paragraph(f"二、样本 {sample['sample_no']} 详细分析", h1_style))
        
        story.append(Paragraph("2.1 基本信息", h2_style))
        
        info_items = [
            ("样本编号", sample.get('sample_no', '')),
            ("采样点", sample.get('sampling_site', '')),
            ("喷发层位", sample.get('eruption_layer', '未记录')),
            ("分组", sample.get('group_name', '未分组')),
            ("采样时间", sample.get('sampling_time', '未记录')),
            ("深度", f"{sample.get('depth', '未记录')} m" if sample.get('depth') else '未记录'),
            ("剖面位置", sample.get('profile_position', '未记录')),
            ("样本总重量", f"{sample.get('total_weight', 0):.2f} g"),
            ("筛分总重量", f"{analysis.get('total_retained', 0):.2f} g"),
            ("底盘重量", f"{analysis.get('pan_weight', 0):.4f} g"),
            ("创建时间", str(sample.get('created_at', ''))),
            ("备注", sample.get('description', '无')),
        ]
        
        info_table_data = []
        for key, value in info_items:
            info_table_data.append([key, str(value)])
        
        info_table = Table(info_table_data, colWidths=[4*cm, 10*cm])
        info_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 10))
        
        story.append(Paragraph("2.2 粒径参数", h2_style))
        
        param_items = [
            ("D10", analysis.get('d10'), 'mm'),
            ("D25", analysis.get('d25'), 'mm'),
            ("中值粒径 D50", analysis.get('median_diameter'), 'mm'),
            ("D75", analysis.get('d75'), 'mm'),
            ("D90", analysis.get('d90'), 'mm'),
            ("分选系数", analysis.get('sorting_coefficient'), ''),
            ("沉积分选评价", analysis.get('sorting_grade', '无法判断'), ''),
            ("可信度评分", f"{analysis.get('reliability_score', 0)} / 100", ''),
            ("可信度等级", analysis.get('reliability_level', '未知'), ''),
        ]
        
        param_table_data = [['参数', '数值']]
        for name, val, unit in param_items:
            if isinstance(val, float):
                display_val = f"{val:.3f} {unit}" if unit else f"{val:.3f}"
            else:
                display_val = str(val) if val else '—'
            param_table_data.append([name, display_val])
        
        param_table = Table(param_table_data, colWidths=[5*cm, 5*cm], repeatRows=1)
        param_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(param_table)
        story.append(Spacer(1, 10))
        
        story.append(Paragraph("2.3 筛分数据明细", h2_style))
        
        df = analysis.get('df')
        if df is not None and not df.empty:
            detail_data = [['粒级', '筛孔 (mm)', '留存重量 (g)', '占比 (%)']]
            
            for _, row_data in df.iterrows():
                detail_data.append([
                    get_sieve_label(row_data['sieve_size']),
                    f"{row_data['sieve_size']}",
                    f"{row_data['retained_weight']:.4f}",
                    f"{row_data['weight_percent']:.2f}",
                ])
            
            pan_w = analysis.get('pan_weight', 0)
            pan_p = analysis.get('pan_percent', 0)
            if pan_w > 0:
                detail_data.append(['底盘（小于最小粒级）', '—', f"{pan_w:.4f}", f"{pan_p:.2f}"])
            
            detail_table = Table(detail_data, colWidths=[4*cm, 3*cm, 3.5*cm, 3.5*cm], repeatRows=1)
            detail_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ]))
            story.append(detail_table)
        
        if idx < len(samples_data):
            story.append(PageBreak())
    
    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def get_import_template_csv() -> str:
    """
    获取批量导入 CSV 模板内容
    """
    headers = [
        '样本编号', '采样点', '喷发层位', '总重量(g)', '备注',
        '分组', '采样时间', '深度(m)', '剖面位置', '纬度', '经度',
    ]
    
    default_sizes = [63.0, 31.5, 16.0, 8.0, 4.0, 2.0, 1.0, 0.5, 0.25, 0.125, 0.063]
    for size in default_sizes:
        headers.append(f'粒级_{size}mm_重量(g)')
    
    example_row = [
        'VA-001', '长白山天池', '全新世喷发层', '100.0', '测试样本',
        '2024春季采集', '2024-03-15', '0.5', '剖面A', '42.02', '128.05',
    ] + ['10.0'] * len(default_sizes)
    
    lines = [
        '# 火山灰样本批量导入模板',
        '# 说明：请在下方填写数据，粒级列填写对应筛孔的留存重量（g）',
        '# 带 * 的为必填项：样本编号、采样点、总重量',
        '# 采样时间格式：YYYY-MM-DD 或 YYYY-MM-DD HH:MM:SS',
        '# 深度单位：米（m）',
        '',
        ','.join(headers),
        ','.join(example_row),
    ]
    
    return '\n'.join(lines)


def parse_import_csv(file_content: str) -> List[Dict]:
    """
    解析批量导入 CSV 文件
    返回样本数据列表
    """
    import csv
    
    lines = []
    for line in file_content.split('\n'):
        line = line.strip()
        if line and not line.startswith('#'):
            lines.append(line)
    
    if len(lines) < 2:
        return []
    
    reader = csv.DictReader(lines)
    
    samples = []
    for row in reader:
        sample = {}
        
        col_mapping = {
            '样本编号': 'sample_no',
            '采样点': 'sampling_site',
            '喷发层位': 'eruption_layer',
            '总重量(g)': 'total_weight',
            '总重量': 'total_weight',
            '备注': 'description',
            '分组': 'group_name',
            '采样时间': 'sampling_time',
            '深度(m)': 'depth',
            '深度': 'depth',
            '剖面位置': 'profile_position',
            '纬度': 'latitude',
            '经度': 'longitude',
        }
        
        for cn_col, en_col in col_mapping.items():
            if cn_col in row:
                sample[en_col] = row[cn_col]
        
        sieve_data = []
        for col_name, value in row.items():
            if col_name.startswith('粒级_') and '_重量' in col_name:
                try:
                    size_str = col_name.replace('粒级_', '').split('_')[0].replace('mm', '')
                    sieve_size = float(size_str)
                    weight = float(value) if value and value.strip() else 0
                    if weight > 0:
                        sieve_data.append({'sieve_size': sieve_size, 'retained_weight': weight})
                except (ValueError, IndexError):
                    continue
        
        sample['sieve_data'] = sieve_data
        samples.append(sample)
    
    return samples
